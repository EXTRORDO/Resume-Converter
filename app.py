from flask import Flask, request, jsonify, send_file, render_template_string
import os
import io
import tempfile
from werkzeug.utils import secure_filename

# Document processing libraries
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import PyPDF2
import pdfplumber
from docx import Document
import mammoth

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

class DocumentConverter:
    @staticmethod
    def text_to_pdf(text_content):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        for paragraph in text_content.split('\n'):
            if paragraph.strip():
                p = Paragraph(paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), styles['Normal'])
                story.append(p)
        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def text_to_docx(text_content):
        doc = Document()
        for paragraph in text_content.split('\n'):
            doc.add_paragraph(paragraph)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def pdf_to_text(pdf_path):
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        return text

    @staticmethod
    def docx_to_text(docx_path):
        try:
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        except:
            doc = Document(docx_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def pdf_to_docx(pdf_path):
        text = DocumentConverter.pdf_to_text(pdf_path)
        return DocumentConverter.text_to_docx(text)

    @staticmethod
    def docx_to_pdf(docx_path):
        text = DocumentConverter.docx_to_text(docx_path)
        return DocumentConverter.text_to_pdf(text)

# HTML Template
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Document Converter</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }
        .container { background: #f9f9f9; padding: 30px; border-radius: 10px; }
        .conversion-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }
        .conversion-btn { padding: 15px; border: 2px solid #ddd; border-radius: 8px; cursor: pointer; text-align: center; transition: all 0.3s; }
        .conversion-btn:hover { border-color: #007bff; background: #f0f8ff; }
        .conversion-btn.active { border-color: #007bff; background: #007bff; color: white; }
        textarea { width: 100%; height: 150px; padding: 10px; border: 1px solid #ddd; border-radius: 5px; }
        .file-upload { border: 2px dashed #ddd; padding: 40px; text-align: center; border-radius: 8px; margin: 20px 0; cursor: pointer; }
        .file-upload:hover { border-color: #007bff; }
        .btn { background: #007bff; color: white; padding: 12px 30px; border: none; border-radius: 5px; cursor: pointer; font-size: 16px; }
        .btn:hover { background: #0056b3; }
        .hidden { display: none; }
        .status { padding: 15px; margin: 20px 0; border-radius: 5px; }
        .success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🔄 Document Converter</h1>
        <p>Convert between Text, PDF, and DOCX formats</p>
        
        <div class="conversion-grid">
            <div class="conversion-btn active" data-type="text-to-pdf">Text → PDF</div>
            <div class="conversion-btn" data-type="text-to-docx">Text → DOCX</div>
            <div class="conversion-btn" data-type="pdf-to-text">PDF → Text</div>
            <div class="conversion-btn" data-type="docx-to-text">DOCX → Text</div>
            <div class="conversion-btn" data-type="pdf-to-docx">PDF → DOCX</div>
            <div class="conversion-btn" data-type="docx-to-pdf">DOCX → PDF</div>
        </div>
        
        <div id="textInput">
            <h3>Enter Text:</h3>
            <textarea id="textContent" placeholder="Enter your text here..."></textarea>
        </div>
        
        <div id="fileInput" class="hidden">
            <h3>Upload File:</h3>
            <div class="file-upload" onclick="document.getElementById('fileUpload').click()">
                <p>Click to select file or drag and drop</p>
                <input type="file" id="fileUpload" style="display:none" accept=".pdf,.docx">
            </div>
            <div id="fileInfo"></div>
        </div>
        
        <button class="btn" onclick="convertDocument()">Convert</button>
        
        <div id="status"></div>
    </div>

    <script>
        let selectedType = 'text-to-pdf';
        let selectedFile = null;
        
        document.querySelectorAll('.conversion-btn').forEach(btn => {
            btn.addEventListener('click', () => {
                document.querySelectorAll('.conversion-btn').forEach(b => b.classList.remove('active'));
                btn.classList.add('active');
                selectedType = btn.dataset.type;
                
                if (selectedType.startsWith('text-')) {
                    document.getElementById('textInput').classList.remove('hidden');
                    document.getElementById('fileInput').classList.add('hidden');
                } else {
                    document.getElementById('textInput').classList.add('hidden');
                    document.getElementById('fileInput').classList.remove('hidden');
                }
            });
        });
        
        document.getElementById('fileUpload').addEventListener('change', (e) => {
            selectedFile = e.target.files[0];
            if (selectedFile) {
                document.getElementById('fileInfo').innerHTML = 
                    `<p><strong>Selected:</strong> ${selectedFile.name} (${(selectedFile.size/1024/1024).toFixed(2)} MB)</p>`;
            }
        });
        
        async function convertDocument() {
            const formData = new FormData();
            formData.append('conversion_type', selectedType);
            
            if (selectedType.startsWith('text-')) {
                const textContent = document.getElementById('textContent').value;
                if (!textContent.trim()) {
                    showStatus('Please enter some text', 'error');
                    return;
                }
                formData.append('content', textContent);
            } else {
                if (!selectedFile) {
                    showStatus('Please select a file', 'error');
                    return;
                }
                formData.append('file', selectedFile);
            }
            
            showStatus('Converting...', 'info');
            
            try {
                const response = await fetch('/api/convert', {
                    method: 'POST',
                    body: formData
                });
                
                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = getFileName();
                    a.click();
                    URL.revokeObjectURL(url);
                    showStatus('Conversion completed successfully!', 'success');
                } else {
                    const error = await response.json();
                    showStatus('Error: ' + error.error, 'error');
                }
            } catch (error) {
                showStatus('Error: ' + error.message, 'error');
            }
        }
        
        function getFileName() {
            const extensions = {
                'text-to-pdf': 'converted.pdf',
                'text-to-docx': 'converted.docx',
                'pdf-to-text': 'extracted.txt',
                'docx-to-text': 'extracted.txt',
                'pdf-to-docx': 'converted.docx',
                'docx-to-pdf': 'converted.pdf'
            };
            return extensions[selectedType];
        }
        
        function showStatus(message, type) {
            const status = document.getElementById('status');
            status.innerHTML = message;
            status.className = 'status ' + type;
        }
    </script>
</body>
</html>
'''

# Route for web interface
@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/api/convert', methods=['POST'])
def convert():
    try:
        conversion_type = request.form.get('conversion_type')
        converter = DocumentConverter()
        
        if conversion_type == 'text-to-pdf':
            content = request.form.get('content')
            buffer = converter.text_to_pdf(content)
            return send_file(buffer, as_attachment=True, download_name='converted.pdf', mimetype='application/pdf')
            
        elif conversion_type == 'text-to-docx':
            content = request.form.get('content')
            buffer = converter.text_to_docx(content)
            return send_file(buffer, as_attachment=True, download_name='converted.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            
        elif conversion_type == 'pdf-to-text':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                text = converter.pdf_to_text(temp_path)
                buffer = io.BytesIO(text.encode('utf-8'))
                return send_file(buffer, as_attachment=True, download_name='extracted.txt', mimetype='text/plain')
            finally:
                os.remove(temp_path)
                    
        elif conversion_type == 'docx-to-text':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                text = converter.docx_to_text(temp_path)
                buffer = io.BytesIO(text.encode('utf-8'))
                return send_file(buffer, as_attachment=True, download_name='extracted.txt', mimetype='text/plain')
            finally:
                os.remove(temp_path)
                    
        elif conversion_type == 'pdf-to-docx':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                buffer = converter.pdf_to_docx(temp_path)
                return send_file(buffer, as_attachment=True, download_name='converted.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            finally:
                os.remove(temp_path)
                    
        elif conversion_type == 'docx-to-pdf':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                buffer = converter.docx_to_pdf(temp_path)
                return send_file(buffer, as_attachment=True, download_name='converted.pdf', mimetype='application/pdf')
            finally:
                os.remove(temp_path)
        
        else:
            return jsonify({'error': 'Invalid conversion type'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("🚀 Document Converter starting...")
    print("📍 Web Interface: http://localhost:5000")
    print("🔗 API Endpoint: http://localhost:5000/api/convert")
    app.run(debug=True, host='0.0.0.0', port=5000)
