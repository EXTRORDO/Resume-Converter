from flask import Flask, request, jsonify, send_file, render_template_string
import os
import io
import tempfile
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import PyPDF2
import pdfplumber
from docx import Document
import mammoth

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

class DocumentConverter:
    @staticmethod
    def text_to_pdf(text_content):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        for paragraph in text_content.split('\n'):
            if paragraph.strip():
                p = Paragraph(paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), styles['Normal'])
                story.append(p)
        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def text_to_docx(text_content):
        doc = Document()
        for paragraph in text_content.split('\n'):
            doc.add_paragraph(paragraph)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def pdf_to_text(pdf_path):
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        return text

    @staticmethod
    def docx_to_text(docx_path):
        try:
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        except:
            doc = Document(docx_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def pdf_to_docx(pdf_path):
        text = DocumentConverter.pdf_to_text(pdf_path)
        return DocumentConverter.text_to_docx(text)

    @staticmethod
    def docx_to_pdf(docx_path):
        text = DocumentConverter.docx_to_text(docx_path)
        return DocumentConverter.text_to_pdf(text)

@app.route('/')
def index():
    return jsonify({"message": "Document Converter API is running"})

@app.route('/api/convert', methods=['POST'])
def convert():
    try:
        conversion_type = request.form.get('conversion_type')
        converter = DocumentConverter()

        if conversion_type == 'text-to-pdf':
            content = request.form.get('content')
            buffer = converter.text_to_pdf(content)
            return send_file(buffer, as_attachment=True, download_name='converted.pdf', mimetype='application/pdf')

        elif conversion_type == 'text-to-docx':
            content = request.form.get('content')
            buffer = converter.text_to_docx(content)
            return send_file(buffer, as_attachment=True, download_name='converted.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        elif conversion_type == 'pdf-to-text':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                text = converter.pdf_to_text(temp_path)
                buffer = io.BytesIO(text.encode('utf-8'))
                return send_file(buffer, as_attachment=True, download_name='extracted.txt', mimetype='text/plain')
            finally:
                os.remove(temp_path)

        elif conversion_type == 'docx-to-text':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                text = converter.docx_to_text(temp_path)
                buffer = io.BytesIO(text.encode('utf-8'))
                return send_file(buffer, as_attachment=True, download_name='extracted.txt', mimetype='text/plain')
            finally:
                os.remove(temp_path)

        elif conversion_type == 'pdf-to-docx':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                buffer = converter.pdf_to_docx(temp_path)
                return send_file(buffer, as_attachment=True, download_name='converted.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            finally:
                os.remove(temp_path)

        elif conversion_type == 'docx-to-pdf':
            file = request.files['file']
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(file.filename))
            file.save(temp_path)
            try:
                buffer = converter.docx_to_pdf(temp_path)
                return send_file(buffer, as_attachment=True, download_name='converted.pdf', mimetype='application/pdf')
            finally:
                os.remove(temp_path)

        else:
            return jsonify({'error': 'Invalid conversion type'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
